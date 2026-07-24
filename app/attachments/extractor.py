"""Trích xuất text từ file theo loại - tương tự app/knowledge/executor.py:
không raise ra ngoài, luôn trả về ExtractionResult (ok/error) để một file
lỗi không bao giờ làm crash luồng upload hay chat."""
from __future__ import annotations

from dataclasses import dataclass, field

from app.attachments.ocr_client import OCRClient, OCRServiceError
from app.core.logging import get_logger

logger = get_logger(__name__)


@dataclass
class ExtractionResult:
    ok: bool
    text: str = ""
    raw: dict = field(default_factory=dict)
    error: str | None = None


def _render_boxes_as_text(boxes: list[dict]) -> str:
    """OCR trả bbox theo vị trí trên ảnh, không theo thứ tự đọc chuẩn -
    sắp theo (y rồi x) để gần giống thứ tự đọc tự nhiên top-to-bottom,
    left-to-right, thay vì giữ nguyên thứ tự trả về."""
    sorted_boxes = sorted(boxes, key=lambda b: (b["bbox"][1], b["bbox"][0]))
    return "\n".join(b["text"] for b in sorted_boxes if b.get("text"))


class AttachmentExtractor:
    def __init__(self) -> None:
        self.ocr_client = OCRClient()

    async def extract_image(self, file_bytes: bytes, *, filename: str, mime_type: str = "image/png") -> ExtractionResult:
        try:
            raw = await self.ocr_client.ocr_image(file_bytes, filename=filename, content_type=mime_type)
            text = _render_boxes_as_text(raw.get("boxes", []))
            if not text.strip():
                return ExtractionResult(ok=False, error="OCR trả về rỗng - không đọc được nội dung ảnh")
            return ExtractionResult(ok=True, text=text, raw=raw)
        except OCRServiceError as exc:
            return ExtractionResult(ok=False, error=str(exc))
    async def extract_pdf(self, file_bytes: bytes, *, filename: str) -> ExtractionResult:
        try:
            raw = await self.ocr_client.ocr_pdf(file_bytes, filename=filename)
            # /ocr/pdf trả về list các trang, mỗi trang có "boxes"
            pages = raw if isinstance(raw, list) else [raw]
            page_texts = []
            for i, page in enumerate(pages, start=1):
                page_text = _render_boxes_as_text(page.get("boxes", []))
                page_texts.append(f"--- Trang {i} ---\n{page_text}")
            return ExtractionResult(ok=True, text="\n\n".join(page_texts), raw={"pages": pages})
        except OCRServiceError as exc:
            return ExtractionResult(ok=False, error=str(exc))

    async def extract_docx(self, file_bytes: bytes, *, filename: str) -> ExtractionResult:
        try:
            import io
            from docx import Document  # python-docx

            doc = Document(io.BytesIO(file_bytes))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            return ExtractionResult(ok=True, text="\n".join(parts))
        except Exception as exc:  # noqa: BLE001
            logger.warning("docx_extraction_failed", filename=filename, error=str(exc))
            return ExtractionResult(ok=False, error=str(exc))

    async def extract_xlsx(self, file_bytes: bytes, *, filename: str) -> ExtractionResult:
        try:
            import io
            import openpyxl

            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
            parts = []
            for sheet in wb.worksheets:
                parts.append(f"--- Sheet: {sheet.title} ---")
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        parts.append(" | ".join(str(c) if c is not None else "" for c in row))
            return ExtractionResult(ok=True, text="\n".join(parts))
        except Exception as exc:  # noqa: BLE001
            logger.warning("xlsx_extraction_failed", filename=filename, error=str(exc))
            return ExtractionResult(ok=False, error=str(exc))